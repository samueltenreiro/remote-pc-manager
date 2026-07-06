import asyncio
import os
import sys
import subprocess
import platform
import time
import json
import requests
from datetime import datetime, timedelta
from pathlib import Path
import threading
import math

from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes

import pyautogui
import psutil
import win32gui
import win32con
import win32process

from PIL import Image, ImageGrab
import cv2

try:
    from docx import Document
    from docx.shared import Pt
except:
    pass

try:
    import PyPDF2
except:
    pass

import requests
import urllib.request
import shutil
import zipfile
from io import BytesIO
import numpy as np

try:
    import pyttsx3
except:
    pass

try:
    import pytesseract
except:
    pass

try:
    from plyer import notification
except:
    pass

try:
    import schedule
except:
    pass

TELEGRAM_TOKEN = "YOUR_TELEGRAM_BOT_TOKEN"
AUTHORIZED_USER = None  

DATA_FILE = "telegram_bot_data.json"

class DataStorage:
    def __init__(self):
        self.data = self.load_data()
    
    def load_data(self):
        if os.path.exists(DATA_FILE):
            with open(DATA_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if 'wifi_state' not in data:
                    data['wifi_state'] = {}
                return data
        return {
            'reminders': [],
            'habits': {},
            'notes': [],
            'authorized_user': None,
            'timers': {},
            'wifi_state': {}
        }
    
    def save_data(self):
        with open(DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(self.data, f, ensure_ascii=False, indent=2)
    
    def add_reminder(self, text, timestamp):
        self.data['reminders'].append({
            'text': text,
            'time': timestamp,
            'notified_10': False,
            'notified_5': False,
            'notified_now': False
        })
        self.save_data()
    
    def get_reminders(self):
        return self.data['reminders']
    
    def remove_reminder(self, index):
        if 0 <= index < len(self.data['reminders']):
            self.data['reminders'].pop(index)
            self.save_data()
            return True
        return False

storage = DataStorage()

class WiFiControl:
    @staticmethod
    def get_current_network():
        """Obtém a rede Wi-Fi atual"""
        try:
            if platform.system() != "Windows":
                return "❌ Só funciona no Windows"
            
            resultado = subprocess.check_output('netsh wlan show interfaces', shell=True, encoding='latin1')
            for linha in resultado.splitlines():
                if "SSID" in linha and "BSSID" not in linha:
                    ssid = linha.split(":", 1)[1].strip()
                    if ssid:
                        return f"📶 **Rede atual:**\n{ssid}"
            return "❌ Não estás ligado a nenhuma rede Wi-Fi"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def list_networks():
        """Lista redes Wi-Fi disponíveis"""
        try:
            if platform.system() != "Windows":
                return "❌ Só funciona no Windows"
            
            resultado = subprocess.check_output('netsh wlan show networks mode=bssid', shell=True, encoding='latin1', timeout=15)
            
            lines = resultado.strip().split('\n')
            networks = []
            current_ssid = None
            
            for line in lines:
                if 'SSID' in line and 'BSSID' not in line:
                    ssid = line.split(':', 1)[1].strip()
                    if ssid:
                        current_ssid = ssid
                elif 'Signal' in line and current_ssid:
                    signal = line.split(':', 1)[1].strip()
                    networks.append(f"• {current_ssid} - {signal}")
                    current_ssid = None
            
            if networks:
                return "📡 **Redes Wi-Fi disponíveis:**\n\n" + "\n".join(networks[:15])
            else:
                return "📡 A processar...\n```\n" + resultado[:500] + "\n```"
        except subprocess.TimeoutExpired:
            return "⏱️ Timeout - a pesquisa demorou muito. Tenta novamente."
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def connect_to_network(ssid, password):
        """Liga a uma rede Wi-Fi"""
        try:
            if platform.system() != "Windows":
                return "❌ Só funciona no Windows"
            
            xml_content = f'''<?xml version="1.0"?>
<WLANProfile xmlns="http://www.microsoft.com/networking/WLAN/profile/v1">
    <name>{ssid}</name>
    <SSIDConfig>
        <SSID>
            <name>{ssid}</name>
        </SSID>
    </SSIDConfig>
    <connectionType>ESS</connectionType>
    <connectionMode>auto</connectionMode>
    <MSM>
        <security>
            <authEncryption>
                <authentication>WPA2PSK</authentication>
                <encryption>AES</encryption>
                <useOneX>false</useOneX>
            </authEncryption>
            <sharedKey>
                <keyType>passPhrase</keyType>
                <protected>false</protected>
                <keyMaterial>{password}</keyMaterial>
            </sharedKey>
        </security>
    </MSM>
</WLANProfile>'''
            
            with open("wifi_temp.xml", "w", encoding="utf-8") as f:
                f.write(xml_content)

            subprocess.run('netsh wlan add profile filename="wifi_temp.xml"', 
                         shell=True, capture_output=True, timeout=10)
            
            os.remove("wifi_temp.xml")

            connect = subprocess.run(f'netsh wlan connect name="{ssid}"', 
                                   shell=True, capture_output=True, text=True, timeout=20)
            
            if "successfully" in connect.stdout.lower() or connect.returncode == 0:
                return f"✅ **LIGADO COM SUCESSO!**\n📶 Rede: {ssid}\n\n✨ O PC já está ligado à nova rede!"
            else:
                return f"⚠️ Perfil criado, mas a ligação pode ter falhado.\nVerifica se:\n• A rede está perto\n• A password está correta\n• O Wi-Fi está ativado"
        except subprocess.TimeoutExpired:
            return "⏱️ Timeout - a ligação demorou muito."
        except Exception as e:
            return f"❌ Erro ao ligar: {str(e)}"
    
    @staticmethod
    def disconnect_wifi():
        """Desliga do Wi-Fi"""
        try:
            if platform.system() != "Windows":
                return "❌ Só funciona no Windows"
            
            subprocess.run('netsh wlan disconnect', shell=True, capture_output=True)
            return "📴 Desligado do Wi-Fi!"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def forget_network(ssid):
        """Remove perfil de rede guardado"""
        try:
            if platform.system() != "Windows":
                return "❌ Só funciona no Windows"
            
            subprocess.run(f'netsh wlan delete profile name="{ssid}"', 
                         shell=True, capture_output=True)
            return f"🗑️ Perfil '{ssid}' removido!"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def show_saved_networks():
        """Mostra redes guardadas"""
        try:
            if platform.system() != "Windows":
                return "❌ Só funciona no Windows"
            
            resultado = subprocess.check_output('netsh wlan show profiles', 
                                              shell=True, encoding='latin1')
            
            profiles = []
            for linha in resultado.splitlines():
                if "All User Profile" in linha or "Perfil de Todos" in linha:
                    profile = linha.split(":", 1)[1].strip()
                    profiles.append(f"• {profile}")
            
            if profiles:
                return "💾 **Redes Wi-Fi guardadas:**\n\n" + "\n".join(profiles)
            else:
                return "📭 Sem redes guardadas"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class SystemControl:
    @staticmethod
    def open_application(app_name):
        """Abre aplicação pelo nome"""
        try:
            if platform.system() == "Windows":
                apps = {
                    'chrome': r'C:\Program Files\Google\Chrome\Application\chrome.exe',
                    'firefox': r'C:\Program Files\Mozilla Firefox\firefox.exe',
                    'edge': r'C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe',
                    'notepad': 'notepad.exe',
                    'calculadora': 'calc.exe',
                    'calc': 'calc.exe',
                    'word': r'C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE',
                    'excel': r'C:\Program Files\Microsoft Office\root\Office16\EXCEL.EXE',
                    'powerpoint': r'C:\Program Files\Microsoft Office\root\Office16\POWERPNT.EXE',
                    'paint': 'mspaint.exe',
                    'explorer': 'explorer.exe',
                    'steam': r'C:\Program Files (x86)\Steam\steam.exe',
                    'discord': os.path.join(os.getenv('LOCALAPPDATA'), 'Discord', 'Update.exe --processStart Discord.exe'),
                    'spotify': os.path.join(os.getenv('APPDATA'), 'Spotify', 'Spotify.exe'),
                    'vscode': os.path.join(os.getenv('LOCALAPPDATA'), 'Programs', 'Microsoft VS Code', 'Code.exe'),
                    'code': os.path.join(os.getenv('LOCALAPPDATA'), 'Programs', 'Microsoft VS Code', 'Code.exe'),
                }
                
                app_lower = app_name.lower()
                
                if app_lower in apps:
                    path = apps[app_lower]
                    if os.path.exists(path):
                        subprocess.Popen([path])
                        return f"✅ {app_name} aberto!"
                    else:
                        alt_path = path.replace('Program Files (x86)', 'Program Files')
                        if os.path.exists(alt_path):
                            subprocess.Popen([alt_path])
                            return f"✅ {app_name} aberto!"
                
                program_files = [
                    r'C:\Program Files',
                    r'C:\Program Files (x86)',
                    os.path.join(os.getenv('LOCALAPPDATA'), 'Programs')
                ]
                
                for base_path in program_files:
                    if os.path.exists(base_path):
                        for root, dirs, files in os.walk(base_path):
                            depth = root[len(base_path):].count(os.sep)
                            if depth > 2:
                                continue
                            
                            for file in files:
                                if app_lower in file.lower() and file.endswith('.exe'):
                                    full_path = os.path.join(root, file)
                                    subprocess.Popen([full_path])
                                    return f"✅ {app_name} aberto! ({file})"
                
                try:
                    subprocess.Popen(app_name, shell=True)
                    return f"✅ Tentei abrir {app_name}"
                except:
                    return f"❌ {app_name} não encontrado. Tenta usar o caminho completo."
            else:
                subprocess.Popen([app_name])
                return f"✅ {app_name} aberto!"
        except Exception as e:
            return f"❌ Erro ao abrir {app_name}: {str(e)}"
    
    @staticmethod
    def close_application(app_name):
        """Fecha aplicação pelo nome"""
        try:
            closed = False
            for proc in psutil.process_iter(['name', 'pid']):
                if app_name.lower() in proc.info['name'].lower():
                    proc.kill()
                    closed = True
            
            if closed:
                return f"✅ {app_name} fechado!"
            else:
                return f"⚠️ {app_name} não está a correr"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def list_running_apps():
        """Lista aplicações em execução"""
        try:
            apps = set()
            
            def enum_windows_callback(hwnd, results):
                if win32gui.IsWindowVisible(hwnd):
                    window_text = win32gui.GetWindowText(hwnd)
                    if window_text and len(window_text) > 1:
                        try:
                            _, pid = win32process.GetWindowThreadProcessId(hwnd)
                            process = psutil.Process(pid)
                            app_name = process.name()
                            
                            system_processes = [
                                'dwm.exe', 'explorer.exe', 'SearchHost.exe', 
                                'StartMenuExperienceHost.exe', 'ShellExperienceHost.exe',
                                'TextInputHost.exe', 'ApplicationFrameHost.exe',
                                'SystemSettings.exe', 'svchost.exe', 'taskhostw.exe'
                            ]
                            
                            if app_name.lower() not in system_processes:
                                apps.add(f"{window_text} ({app_name})")
                        except:
                            pass
                return True
            
            win32gui.EnumWindows(enum_windows_callback, None)
            
            if apps:
                return "🖥️ **Aplicações abertas:**\n" + "\n".join(f"• {app}" for app in sorted(apps))
            else:
                return "⚠️ Nenhuma aplicação visível"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def shutdown_pc():
        """Desliga o PC"""
        if platform.system() == "Windows":
            os.system("shutdown /s /t 5")
        else:
            os.system("shutdown -h now")
        return "🔴 PC vai desligar em 5 segundos!"
    
    @staticmethod
    def restart_pc():
        """Reinicia o PC"""
        if platform.system() == "Windows":
            os.system("shutdown /r /t 5")
        else:
            os.system("shutdown -r now")
        return "🔄 PC vai reiniciar em 5 segundos!"
    
    @staticmethod
    def sleep_pc():
        """Suspende o PC"""
        if platform.system() == "Windows":
            os.system("rundll32.exe powrprof.dll,SetSuspendState 0,1,0")
        else:
            os.system("systemctl suspend")
        return "😴 PC vai suspender!"

class Automation:
    @staticmethod
    def move_mouse(x, y, relative=False):
        """Move o rato"""
        try:
            if relative:
                current_x, current_y = pyautogui.position()
                pyautogui.moveTo(current_x + x, current_y + y, duration=0.2)
            else:
                pyautogui.moveTo(x, y, duration=0.2)
            return f"✅ Rato movido para ({x}, {y})"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def click_mouse(button='left'):
        """Clica o rato"""
        try:
            pyautogui.click(button=button)
            return f"✅ Clique {button} executado!"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def type_text(text):
        """Escreve texto"""
        try:
            pyautogui.write(text, interval=0.05)
            return f"✅ Texto escrito: {text[:50]}..."
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def press_key(key):
        """Pressiona tecla"""
        try:
            pyautogui.press(key)
            return f"✅ Tecla {key} pressionada!"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def hotkey(*keys):
        """Executa combinação de teclas"""
        try:
            pyautogui.hotkey(*keys)
            return f"✅ Atalho {'+'.join(keys)} executado!"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def take_screenshot():
        """Tira screenshot"""
        try:
            screenshot = ImageGrab.grab()
            filename = f"screenshot_{int(time.time())}.png"
            screenshot.save(filename)
            return filename
        except Exception as e:
            return None

class FileManager:
    @staticmethod
    def find_files(query, search_content=False):
        """Procura ficheiros por nome"""
        try:
            results = []
            home = str(Path.home())
            
            for root, dirs, files in os.walk(home):
                if root.count(os.sep) - home.count(os.sep) > 3:
                    continue
                
                for file in files:
                    if query.lower() in file.lower():
                        results.append(os.path.join(root, file))
                    
                    if len(results) >= 20:
                        break
                
                if len(results) >= 20:
                    break
            
            if results:
                return "📁 **Ficheiros encontrados:**\n" + "\n".join(f"• {r}" for r in results[:20])
            else:
                return "❌ Nenhum ficheiro encontrado"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def open_file(filepath):
        """Abre ficheiro"""
        try:
            if os.path.exists(filepath):
                os.startfile(filepath) if platform.system() == "Windows" else subprocess.call(['open', filepath])
                return f"✅ Ficheiro aberto: {filepath}"
            else:
                return "❌ Ficheiro não existe"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class AdvancedFileManager:
    @staticmethod
    def copy_file(source, dest):
        """Copia ficheiro"""
        try:
            shutil.copy2(source, dest)
            return f"✅ Ficheiro copiado de {source} para {dest}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def move_file(source, dest):
        """Move ficheiro"""
        try:
            shutil.move(source, dest)
            return f"✅ Ficheiro movido de {source} para {dest}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def delete_file(filepath):
        """Apaga ficheiro"""
        try:
            if os.path.isfile(filepath):
                os.remove(filepath)
                return f"✅ Ficheiro apagado: {filepath}"
            elif os.path.isdir(filepath):
                shutil.rmtree(filepath)
                return f"✅ Pasta apagada: {filepath}"
            else:
                return "❌ Ficheiro não existe"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def rename_file(old_path, new_name):
        """Renomeia ficheiro"""
        try:
            directory = os.path.dirname(old_path)
            new_path = os.path.join(directory, new_name)
            os.rename(old_path, new_path)
            return f"✅ Ficheiro renomeado para {new_name}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def compress_files(files, zip_name):
        """Comprime ficheiros em ZIP"""
        try:
            with zipfile.ZipFile(zip_name, 'w') as zipf:
                for file in files:
                    if os.path.exists(file):
                        zipf.write(file, os.path.basename(file))
            return f"✅ Ficheiros comprimidos em {zip_name}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def extract_zip(zip_path, extract_to="."):
        """Extrai ficheiro ZIP"""
        try:
            with zipfile.ZipFile(zip_path, 'r') as zipf:
                zipf.extractall(extract_to)
            return f"✅ Ficheiros extraídos para {extract_to}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def get_file_info(filepath):
        """Obtém informação do ficheiro"""
        try:
            if not os.path.exists(filepath):
                return "❌ Ficheiro não existe"
            
            stat = os.stat(filepath)
            size = stat.st_size
            created = datetime.fromtimestamp(stat.st_ctime)
            modified = datetime.fromtimestamp(stat.st_mtime)
            
            size_str = f"{size} bytes"
            if size > 1024**3:
                size_str = f"{size / 1024**3:.2f} GB"
            elif size > 1024**2:
                size_str = f"{size / 1024**2:.2f} MB"
            elif size > 1024:
                size_str = f"{size / 1024:.2f} KB"
            
            return f"""📄 **Informação do ficheiro:**

📝 Nome: {os.path.basename(filepath)}
📏 Tamanho: {size_str}
📅 Criado: {created.strftime('%d/%m/%Y %H:%M')}
✏️ Modificado: {modified.strftime('%d/%m/%Y %H:%M')}
📂 Caminho: {filepath}"""
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class Calculator:
    @staticmethod
    def calculate(expression):
        """Calcula expressão matemática"""
        try:
            allowed_chars = set('0123456789+-*/().,^ ')
            if not all(c in allowed_chars for c in expression):
                return "❌ Expressão inválida"
            
            expression = expression.replace('^', '**')
            result = eval(expression)
            return f"🧮 **Resultado:** {result}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def convert_units(value, from_unit, to_unit):
        """Converte unidades"""
        conversions = {
            ('m', 'cm'): 100,
            ('m', 'km'): 0.001,
            ('cm', 'm'): 0.01,
            ('km', 'm'): 1000,
            ('m', 'ft'): 3.28084,
            ('ft', 'm'): 0.3048,
            ('kg', 'g'): 1000,
            ('g', 'kg'): 0.001,
            ('kg', 'lb'): 2.20462,
            ('lb', 'kg'): 0.453592,
            ('c', 'f'): lambda x: x * 9/5 + 32,
            ('f', 'c'): lambda x: (x - 32) * 5/9,
            ('c', 'k'): lambda x: x + 273.15,
            ('k', 'c'): lambda x: x - 273.15,
        }
        
        key = (from_unit.lower(), to_unit.lower())
        
        if key in conversions:
            factor = conversions[key]
            if callable(factor):
                result = factor(value)
            else:
                result = value * factor
            return f"✅ {value} {from_unit} = {result:.2f} {to_unit}"
        else:
            return "❌ Conversão não suportada"

class Weather:
    @staticmethod
    def get_weather(city):
        """Obtém meteorologia"""
        try:
            url = f"https://wttr.in/{city}?format=j1"
            response = requests.get(url, timeout=5)
            
            if response.status_code == 200:
                data = response.json()
                current = data['current_condition'][0]
                
                weather_pt = {
                    'Sunny': '☀️ Ensolarado',
                    'Clear': '🌙 Limpo',
                    'Partly cloudy': '⛅ Parcialmente nublado',
                    'Cloudy': '☁️ Nublado',
                    'Overcast': '☁️ Muito nublado',
                    'Mist': '🌫️ Névoa',
                    'Fog': '🌫️ Nevoeiro',
                    'Light rain': '🌧️ Chuva leve',
                    'Moderate rain': '🌧️ Chuva moderada',
                    'Heavy rain': '🌧️ Chuva forte',
                    'Light snow': '🌨️ Neve leve',
                    'Heavy snow': '🌨️ Neve forte',
                    'Thunderstorm': '⛈️ Trovoada',
                }
                
                weather_desc = current['weatherDesc'][0]['value']
                weather_emoji = weather_pt.get(weather_desc, f"🌤️ {weather_desc}")
                
                return f"""🌍 **Meteorologia em {city.title()}:**

{weather_emoji}
🌡️ **Temperatura:** {current['temp_C']}°C (sensação: {current['FeelsLikeC']}°C)
💧 **Humidade:** {current['humidity']}%
💨 **Vento:** {current['windspeedKmph']} km/h
👁️ **Visibilidade:** {current['visibility']} km
☁️ **Cobertura de nuvens:** {current['cloudcover']}%
🌧️ **Precipitação:** {current['precipMM']} mm"""
            else:
                return f"❌ Cidade '{city}' não encontrada"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class WebControl:
    @staticmethod
    def google_search(query):
        """Pesquisa no Google"""
        try:
            search_url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
            
            chrome_paths = [
                r'C:\Program Files\Google\Chrome\Application\chrome.exe',
                r'C:\Program Files (x86)\Google\Chrome\Application\chrome.exe',
            ]
            
            chrome_opened = False
            for chrome_path in chrome_paths:
                if os.path.exists(chrome_path):
                    subprocess.Popen([chrome_path, search_url])
                    chrome_opened = True
                    break
            
            if not chrome_opened:
                import webbrowser
                webbrowser.open(search_url)
            return f"📺 A pesquisar no Google: '{query}'"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def open_website(url):
        """Abre website"""
        try:
            import webbrowser
            if not url.startswith('http'):
                url = 'https://' + url
            webbrowser.open(url)
            return f"🌐 A abrir: {url}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class InternetControl:
    @staticmethod
    def test_speed():
        """Testa velocidade da internet"""
        try:
            import speedtest
            st = speedtest.Speedtest()
            st.get_best_server()
            download = st.download() / 1_000_000
            upload = st.upload() / 1_000_000
            ping = st.results.ping
            
            return f"""🌐 **Teste de Velocidade:**

⬇️ **Download:** {download:.2f} Mbps
⬆️ **Upload:** {upload:.2f} Mbps
📶 **Ping:** {ping:.2f} ms"""
        except Exception as e:
            return f"❌ Erro: {str(e)}\nInstala: pip install speedtest-cli"
    
    @staticmethod
    def get_ip_info():
        """Obtém informação de IP"""
        try:
            response = requests.get('https://api.ipify.org?format=json', timeout=5)
            ip = response.json()['ip']
            return f"🌐 **Teu IP Público:** {ip}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def download_file(url, filename=None):
        """Descarrega ficheiro"""
        try:
            if filename is None:
                filename = url.split('/')[-1]
            
            response = requests.get(url, stream=True)
            with open(filename, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            return f"✅ Ficheiro descarregado: {filename}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def check_website(url):
        """Verifica se website está online"""
        try:
            if not url.startswith('http'):
                url = 'https://' + url
            
            response = requests.get(url, timeout=5)
            return f"✅ {url} está ONLINE (Status: {response.status_code})"
        except Exception as e:
            return f"❌ {url} está OFFLINE ou inacessível"

class NotesSystem:
    @staticmethod
    def add_note(title, content):
        """Adiciona nota"""
        storage.data['notes'].append({
            'title': title,
            'content': content,
            'created': datetime.now().isoformat()
        })
        storage.save_data()
        return "✅ Nota guardada!"
    
    @staticmethod
    def list_notes():
        """Lista notas"""
        if not storage.data['notes']:
            return "🔭 Sem notas guardadas"
        
        text = "📝 **Notas guardadas:**\n\n"
        for i, note in enumerate(storage.data['notes']):
            created = datetime.fromisoformat(note['created'])
            text += f"{i+1}. **{note['title']}**\n"
            text += f"   📅 {created.strftime('%d/%m/%Y %H:%M')}\n"
            text += f"   {note['content'][:50]}...\n\n"
        return text
    
    @staticmethod
    def get_note(index):
        """Mostra nota completa"""
        if 0 <= index < len(storage.data['notes']):
            note = storage.data['notes'][index]
            created = datetime.fromisoformat(note['created'])
            return f"📝 **{note['title']}**\n📅 {created.strftime('%d/%m/%Y %H:%M')}\n\n{note['content']}"
        return "❌ Nota não existe"
    
    @staticmethod
    def delete_note(index):
        """Apaga nota"""
        if 0 <= index < len(storage.data['notes']):
            storage.data['notes'].pop(index)
            storage.save_data()
            return "✅ Nota apagada!"
        return "❌ Nota não existe"

class VolumeControl:
    @staticmethod
    def set_volume(level):
        """Define volume"""
        try:
            if platform.system() == "Windows":
                from ctypes import cast, POINTER
                from comtypes import CLSCTX_ALL
                from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
                
                devices = AudioUtilities.GetSpeakers()
                interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
                volume = cast(interface, POINTER(IAudioEndpointVolume))
                volume.SetMasterVolumeLevelScalar(level / 100, None)
                return f"🔊 Volume definido para {level}%"
        except:
            return "❌ Erro. Instala: pip install pycaw comtypes"
    
    @staticmethod
    def mute():
        """Silencia"""
        try:
            Automation.hotkey('volumemute')
            return "🔇 Som silenciado/reativado"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class MediaControl:
    @staticmethod
    def play_pause():
        """Play/Pause"""
        try:
            Automation.hotkey('playpause')
            return "⏯️ Play/Pause"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def next_track():
        """Próxima música"""
        try:
            Automation.hotkey('nexttrack')
            return "⏭️ Próxima música"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def prev_track():
        """Música anterior"""
        try:
            Automation.hotkey('prevtrack')
            return "⏮️ Música anterior"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class SystemInfo:
    @staticmethod
    def get_system_info():
        """Informação do sistema"""
        try:
            cpu_percent = psutil.cpu_percent(interval=1)
            memory = psutil.virtual_memory()
            disk = psutil.disk_usage('/')
            
            return f"""💻 **Informação do Sistema:**

🖥️ **CPU:** {cpu_percent}%
🧠 **RAM:** {memory.percent}% ({memory.used // (1024**3)}GB / {memory.total // (1024**3)}GB)
💾 **Disco:** {disk.percent}% ({disk.used // (1024**3)}GB / {disk.total // (1024**3)}GB)
🔋 **Bateria:** {SystemInfo.get_battery()}

⚙️ **Sistema:** {platform.system()} {platform.release()}
🖥️ **Arquitetura:** {platform.machine()}"""
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def get_battery():
        """Info da bateria"""
        try:
            battery = psutil.sensors_battery()
            if battery:
                plugged = "🔌 Ligado" if battery.power_plugged else "🔋 Bateria"
                return f"{battery.percent}% ({plugged})"
            return "Sem bateria"
        except:
            return "N/A"

class TaskManager:
    @staticmethod
    def kill_process_by_name(name):
        """Mata processo"""
        try:
            killed = 0
            for proc in psutil.process_iter(['name', 'pid']):
                if name.lower() in proc.info['name'].lower():
                    proc.kill()
                    killed += 1
            
            if killed > 0:
                return f"✅ {killed} processo(s) '{name}' terminado(s)"
            return f"⚠️ Processo '{name}' não encontrado"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def list_processes():
        """Lista processos"""
        try:
            processes = []
            for proc in psutil.process_iter(['name', 'cpu_percent', 'memory_percent']):
                try:
                    info = proc.info
                    if info['cpu_percent'] > 0 or info['memory_percent'] > 1:
                        processes.append(info)
                except:
                    pass
            
            processes.sort(key=lambda x: x['cpu_percent'], reverse=True)
            
            text = "📊 **Processos (Top 10 CPU):**\n\n"
            for i, p in enumerate(processes[:10]):
                text += f"{i+1}. **{p['name']}**\n"
                text += f"   CPU: {p['cpu_percent']:.1f}% | RAM: {p['memory_percent']:.1f}%\n\n"
            
            return text
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class WordAutomation:
    @staticmethod
    def create_document(title, content, pages=1, font_size=12):
        """Cria documento Word"""
        try:
            doc = Document()
            doc.add_heading(title, 0)
            
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                p = doc.add_paragraph(para)
                for run in p.runs:
                    run.font.size = Pt(font_size)
            
            filename = f"{title.replace(' ', '_')}_{int(time.time())}.docx"
            doc.save(filename)
            
            return f"✅ Documento criado: {filename}"
        except Exception as e:
            return f"❌ Erro: {str(e)}\nInstala: pip install python-docx"

class CameraControl:
    @staticmethod
    def take_photo():
        """Tira foto da webcam"""
        try:
            cap = cv2.VideoCapture(0)
            ret, frame = cap.read()
            
            if ret:
                filename = f"webcam_{int(time.time())}.jpg"
                cv2.imwrite(filename, frame)
                cap.release()
                return filename
            
            cap.release()
            return None
        except Exception as e:
            return None

class TimerControl:
    @staticmethod
    def start_timer(name, seconds):
        """Inicia timer"""
        end_time = datetime.now() + timedelta(seconds=seconds)
        storage.data['timers'][name] = {
            'end_time': end_time.isoformat(),
            'seconds': seconds
        }
        storage.save_data()
        return f"⏱️ Timer '{name}' iniciado por {seconds}s!"
    
    @staticmethod
    def list_timers():
        """Lista timers"""
        if not storage.data['timers']:
            return "⏱️ Sem timers ativos"
        
        text = "⏱️ **Timers ativos:**\n\n"
        now = datetime.now()
        
        for name, timer in storage.data['timers'].items():
            end_time = datetime.fromisoformat(timer['end_time'])
            remaining = (end_time - now).total_seconds()
            
            if remaining > 0:
                text += f"• **{name}**: {int(remaining)}s restantes\n"
            else:
                text += f"• **{name}**: ⏰ TERMINADO!\n"
        
        return text
    
    @staticmethod
    def check_timer(name):
        """Verifica timer específico"""
        if name not in storage.data['timers']:
            return f"❌ Timer '{name}' não existe"
        
        timer = storage.data['timers'][name]
        end_time = datetime.fromisoformat(timer['end_time'])
        now = datetime.now()
        remaining = (end_time - now).total_seconds()
        
        if remaining > 0:
            return f"⏱️ Timer '{name}': {int(remaining)}s restantes"
        else:
            return f"⏰ Timer '{name}' TERMINADO!"

class SecurityControl:
    @staticmethod
    def lock_pc():
        """Bloqueia PC"""
        try:
            if platform.system() == "Windows":
                os.system("rundll32.exe user32.dll,LockWorkStation")
            else:
                os.system("gnome-screensaver-command -l")
            return "🔒 PC bloqueado!"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def clear_history():
        """Limpa histórico do navegador"""
        return "⚠️ Funcionalidade em desenvolvimento"

class ImageControl:
    @staticmethod
    def resize_image(image_path, width, height):
        """Redimensiona imagem"""
        try:
            img = Image.open(image_path)
            img_resized = img.resize((width, height))
            
            name, ext = os.path.splitext(image_path)
            new_path = f"{name}_resized{ext}"
            img_resized.save(new_path)
            
            return f"✅ Imagem redimensionada: {new_path}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def convert_image(image_path, new_format):
        """Converte formato de imagem"""
        try:
            img = Image.open(image_path)
            
            name = os.path.splitext(image_path)[0]
            new_path = f"{name}.{new_format}"
            img.save(new_path)
            
            return f"✅ Imagem convertida: {new_path}"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class TTSControl:
    @staticmethod
    def speak(text):
        """Fala texto"""
        try:
            engine = pyttsx3.init()
            engine.say(text)
            engine.runAndWait()
            return f"🔊 A falar: '{text}'"
        except Exception as e:
            return f"❌ Erro: {str(e)}\nInstala: pip install pyttsx3"

class NotificationControl:
    @staticmethod
    def show_notification(title, message):
        """Mostra notificação"""
        try:
            notification.notify(
                title=title,
                message=message,
                timeout=10
            )
            return f"✅ Notificação enviada: {title}"
        except Exception as e:
            return f"❌ Erro: {str(e)}\nInstala: pip install plyer"

class MaintenanceControl:
    @staticmethod
    def clean_temp():
        """Limpa ficheiros temporários"""
        try:
            temp_dir = os.environ.get('TEMP', '/tmp')
            deleted = 0
            
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    try:
                        os.remove(os.path.join(root, file))
                        deleted += 1
                    except:
                        pass
            
            return f"✅ {deleted} ficheiros temporários apagados"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def empty_recycle_bin():
        """Esvazia reciclagem"""
        try:
            if platform.system() == "Windows":
                os.system('rd /s /q %systemdrive%\\$Recycle.bin')
            return "✅ Reciclagem esvaziada!"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def check_disk_space():
        """Verifica espaço em disco"""
        try:
            text = "💾 **Espaço em Disco:**\n\n"
            
            for partition in psutil.disk_partitions():
                try:
                    usage = psutil.disk_usage(partition.mountpoint)
                    text += f"**{partition.device}**\n"
                    text += f"  Total: {usage.total // (1024**3)}GB\n"
                    text += f"  Usado: {usage.used // (1024**3)}GB ({usage.percent}%)\n"
                    text += f"  Livre: {usage.free // (1024**3)}GB\n\n"
                except:
                    pass
            
            return text
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class CustomizationControl:
    @staticmethod
    def change_wallpaper(image_path):
        """Muda papel de parede"""
        try:
            if platform.system() == "Windows":
                import ctypes
                ctypes.windll.user32.SystemParametersInfoW(20, 0, image_path, 3)
                return f"✅ Papel de parede alterado!"
            return "❌ Apenas suportado no Windows"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

class GamingControl:
    @staticmethod
    def gaming_mode():
        """Ativa modo gaming"""
        try:
            # Fecha apps desnecessárias
            apps_to_close = ['chrome.exe', 'firefox.exe', 'discord.exe']
            closed = 0
            
            for proc in psutil.process_iter(['name']):
                if proc.info['name'].lower() in apps_to_close:
                    try:
                        proc.kill()
                        closed += 1
                    except:
                        pass
            
            return f"🎮 Modo Gaming ativado! {closed} apps fechadas"
        except Exception as e:
            return f"❌ Erro: {str(e)}"
    
    @staticmethod
    def boost_performance():
        """Otimiza performance"""
        try:
            
            import gc
            gc.collect()
            return "⚡ Performance otimizada!"
        except Exception as e:
            return f"❌ Erro: {str(e)}"

async def check_reminders(app):
    """Verifica lembretes"""
    while True:
        try:
            now = datetime.now()
            reminders = storage.get_reminders()
            
            for i, reminder in enumerate(reminders):
                reminder_time = datetime.fromisoformat(reminder['time'])
                time_diff = (reminder_time - now).total_seconds()
                
                if 595 < time_diff <= 605 and not reminder['notified_10']:
                    await app.bot.send_message(
                        chat_id=storage.data['authorized_user'],
                        text=f"⏰ **Lembrete em 10 minutos:**\n{reminder['text']}"
                    )
                    reminder['notified_10'] = True
                    storage.save_data()
                
                elif 295 < time_diff <= 305 and not reminder['notified_5']:
                    await app.bot.send_message(
                        chat_id=storage.data['authorized_user'],
                        text=f"⏰ **Lembrete em 5 minutos:**\n{reminder['text']}"
                    )
                    reminder['notified_5'] = True
                    storage.save_data()
                
                elif -5 < time_diff <= 5 and not reminder['notified_now']:
                    await app.bot.send_message(
                        chat_id=storage.data['authorized_user'],
                        text=f"🔔 **LEMBRETE AGORA:**\n{reminder['text']}"
                    )
                    reminder['notified_now'] = True
                    storage.save_data()
            
            storage.data['reminders'] = [
                r for r in reminders
                if (datetime.fromisoformat(r['time']) - now).total_seconds() > -3600
            ]
            storage.save_data()
            
        except Exception as e:
            print(f"Erro ao verificar lembretes: {e}")
        
        await asyncio.sleep(30)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Comando /start"""
    chat_id = update.effective_chat.id
    
    if storage.data['authorized_user'] is None:
        storage.data['authorized_user'] = chat_id
        storage.save_data()
    
    await update.message.reply_text(f"""🤖 **Sistema de Controlo Total do PC + Wi-Fi Ativo!**

📱 **Teu Chat ID:** `{chat_id}`

Escreve `ajuda` para ver todos os comandos! 🚀
✨ **Novidade:** Controlo completo de Wi-Fi! 📶""")

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Processa mensagens gerais"""
    text = update.message.text.strip()
    text_lower = text.lower()
    parts = text.split()
    
    if not parts:
        return
    
    comando = parts[0].lower()
    args = parts[1:] if len(parts) > 1 else []

    if comando in ['wifi', 'rede']:
        result = WiFiControl.get_current_network()
        await update.message.reply_text(result)
    
    elif comando in ['redes', 'listar', 'scan']:
        await update.message.reply_text("🔍 A procurar redes Wi-Fi... (5-10 seg)")
        result = WiFiControl.list_networks()
        await update.message.reply_text(result)
    
    elif comando in ['ligar'] and len(args) > 0 and args[0] not in ['pc', 'sistema']:
        user_id = update.effective_user.id
        storage.data['wifi_state'][str(user_id)] = {'step': 'waiting_password', 'ssid': ' '.join(args)}
        storage.save_data()
        await update.message.reply_text(f"🔐 **Nome da rede:** {' '.join(args)}\n\nAgora envia a **password** da rede:")
    
    elif comando in ['desligar'] and len(args) > 0 and args[0] == 'wifi':
        result = WiFiControl.disconnect_wifi()
        await update.message.reply_text(result)
    
    elif comando in ['esquecer', 'remover'] and len(args) > 0:
        result = WiFiControl.forget_network(' '.join(args))
        await update.message.reply_text(result)
    
    elif comando in ['guardadas', 'saved']:
        result = WiFiControl.show_saved_networks()
        await update.message.reply_text(result)

    elif str(update.effective_user.id) in storage.data['wifi_state']:
        state = storage.data['wifi_state'][str(update.effective_user.id)]
        if state['step'] == 'waiting_password':
            ssid = state['ssid']
            password = text.strip()
            
            await update.message.reply_text(f"🔄 A ligar a **{ssid}**...\nAguarda 10-15 segundos...")
            result = WiFiControl.connect_to_network(ssid, password)
            await update.message.reply_text(result)
            
            del storage.data['wifi_state'][str(update.effective_user.id)]
            storage.save_data()

    elif comando in ['abrir', 'abre', 'open']:
        if not args:
            await update.message.reply_text("❌ Uso: abrir [app]\nEx: abrir chrome")
            return
        result = SystemControl.open_application(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['fechar', 'fecha', 'close']:
        if not args:
            await update.message.reply_text("❌ Uso: fechar [app]")
            return
        result = SystemControl.close_application(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['apps', 'aplicações', 'janelas']:
        result = SystemControl.list_running_apps()
        await update.message.reply_text(result)
    
    elif comando in ['desligar', 'shutdown']:
        result = SystemControl.shutdown_pc()
        await update.message.reply_text(result)
    
    elif comando in ['reiniciar', 'restart']:
        result = SystemControl.restart_pc()
        await update.message.reply_text(result)
    
    elif comando in ['suspender', 'sleep']:
        result = SystemControl.sleep_pc()
        await update.message.reply_text(result)

    elif comando in ['rato', 'mouse']:
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: rato [x] [y]\nEx: rato 100 50")
            return
        try:
            x, y = int(args[0]), int(args[1])
            current_x, current_y = pyautogui.position()
            result = Automation.move_mouse(x, y, relative=True)
            await update.message.reply_text(f"{result}\n📍 ({current_x},{current_y}) → ({current_x+x},{current_y+y})")
        except:
            await update.message.reply_text("❌ Valores inválidos")
    
    elif comando in ['click', 'clica']:
        result = Automation.click_mouse()
        await update.message.reply_text(result)
    
    elif comando in ['escrever', 'type']:
        if not args:
            await update.message.reply_text("❌ Uso: escrever [texto]")
            return
        result = Automation.type_text(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['tecla', 'press']:
        if not args:
            await update.message.reply_text("❌ Uso: tecla [nome]")
            return
        result = Automation.press_key(args[0])
        await update.message.reply_text(result)
    
    elif comando in ['atalho', 'hotkey']:
        if not args:
            await update.message.reply_text("❌ Uso: atalho ctrl+c")
            return
        keys = args[0].split('+')
        result = Automation.hotkey(*keys)
        await update.message.reply_text(result)
    
    elif comando in ['screenshot', 'print']:
        await update.message.reply_text("📸 A tirar screenshot...")
        filename = Automation.take_screenshot()
        if filename:
            await update.message.reply_photo(photo=open(filename, 'rb'))
            os.remove(filename)
        else:
            await update.message.reply_text("❌ Erro")
    
    elif comando in ['procurar', 'search', 'find']:
        if not args:
            await update.message.reply_text("❌ Uso: procurar [nome]")
            return
        result = FileManager.find_files(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['copiar', 'copy']:
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: copiar [origem] [destino]")
            return
        result = AdvancedFileManager.copy_file(args[0], " ".join(args[1:]))
        await update.message.reply_text(result)
    
    elif comando in ['mover', 'move']:
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: mover [origem] [destino]")
            return
        result = AdvancedFileManager.move_file(args[0], " ".join(args[1:]))
        await update.message.reply_text(result)
    
    elif comando in ['apagar', 'delete'] and len(args) > 0 and args[0] not in ['nota', 'lembrete']:
        result = AdvancedFileManager.delete_file(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['renomear', 'rename']:
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: renomear [ficheiro] [novo_nome]")
            return
        result = AdvancedFileManager.rename_file(args[0], " ".join(args[1:]))
        await update.message.reply_text(result)
    
    elif comando in ['zip', 'comprimir']:
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: zip [nome.zip] [ficheiro1] [ficheiro2]...")
            return
        result = AdvancedFileManager.compress_files(args[1:], args[0])
        await update.message.reply_text(result)
    
    elif comando in ['unzip', 'extrair']:
        if not args:
            await update.message.reply_text("❌ Uso: unzip [ficheiro.zip]")
            return
        extract_to = args[1] if len(args) > 1 else "."
        result = AdvancedFileManager.extract_zip(args[0], extract_to)
        await update.message.reply_text(result)
    
    elif comando in ['info'] and len(args) > 0 and args[0] not in ['sistema', 'system']:
        result = AdvancedFileManager.get_file_info(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['calc', 'calcular']:
        if not args:
            await update.message.reply_text("❌ Uso: calc [expressão]")
            return
        result = Calculator.calculate(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['converter', 'convert']:
        if len(args) < 3:
            await update.message.reply_text("❌ Uso: converter [valor] [de] [para]")
            return
        try:
            result = Calculator.convert_units(float(args[0]), args[1], args[2])
            await update.message.reply_text(result)
        except:
            await update.message.reply_text("❌ Valores inválidos")

    elif comando in ['lembrete', 'reminder']:
        if len(args) < 3:
            await update.message.reply_text("❌ Uso: lembrete [texto] YYYY-MM-DD HH:MM")
            return
        try:
            text_parts = args[:-2]
            reminder_text = " ".join(text_parts)
            datetime_str = f"{args[-2]} {args[-1]}"
            reminder_time = datetime.strptime(datetime_str, "%Y-%m-%d %H:%M")
            storage.add_reminder(reminder_text, reminder_time.isoformat())
            await update.message.reply_text(f"✅ Lembrete criado para {reminder_time.strftime('%d/%m/%Y %H:%M')}!")
        except:
            await update.message.reply_text("❌ Formato inválido")
    
    elif comando in ['lembretes', 'reminders']:
        reminders = storage.get_reminders()
        if not reminders:
            await update.message.reply_text("🔭 Sem lembretes")
            return
        text_response = "⏰ **Lembretes:**\n\n"
        for i, r in enumerate(reminders):
            time = datetime.fromisoformat(r['time'])
            text_response += f"{i+1}. {r['text']}\n   📅 {time.strftime('%d/%m/%Y %H:%M')}\n\n"
        await update.message.reply_text(text_response)
    
    elif comando in ['apagar'] and len(args) > 0 and args[0] == 'lembrete':
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: apagar lembrete [número]")
            return
        try:
            if storage.remove_reminder(int(args[1]) - 1):
                await update.message.reply_text("✅ Lembrete apagado!")
            else:
                await update.message.reply_text("❌ Não existe")
        except:
            await update.message.reply_text("❌ Número inválido")

    elif comando in ['google', 'pesquisar']:
        if not args:
            await update.message.reply_text("❌ Uso: google [pesquisa]")
            return
        result = WebControl.google_search(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['youtube', 'yt']:
        if not args:
            await update.message.reply_text("❌ Uso: youtube [pesquisa]")
            return
        result = WebControl.youtube_search(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['site', 'website']:
        if not args:
            await update.message.reply_text("❌ Uso: site [url]")
            return
        result = WebControl.open_website(" ".join(args))
        await update.message.reply_text(result)

    elif comando in ['nota', 'note']:
        if not args:
            await update.message.reply_text("❌ Uso: nota [título] | [texto]")
            return
        full_text = " ".join(args)
        if '|' not in full_text:
            await update.message.reply_text("❌ Use | para separar")
            return
        parts = full_text.split('|', 1)
        result = NotesSystem.add_note(parts[0].strip(), parts[1].strip())
        await update.message.reply_text(result)
    
    elif comando in ['notas', 'notes']:
        result = NotesSystem.list_notes()
        await update.message.reply_text(result)
    
    elif comando in ['ver'] and len(args) > 0 and args[0] == 'nota':
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: ver nota [número]")
            return
        try:
            result = NotesSystem.get_note(int(args[1]) - 1)
            await update.message.reply_text(result)
        except:
            await update.message.reply_text("❌ Número inválido")
    
    elif comando in ['apagar'] and len(args) > 0 and args[0] == 'nota':
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: apagar nota [número]")
            return
        try:
            result = NotesSystem.delete_note(int(args[1]) - 1)
            await update.message.reply_text(result)
        except:
            await update.message.reply_text("❌ Número inválido")

    elif comando in ['volume', 'vol']:
        if not args:
            await update.message.reply_text("❌ Uso: volume [0-100]")
            return
        try:
            level = int(args[0])
            if 0 <= level <= 100:
                result = VolumeControl.set_volume(level)
                await update.message.reply_text(result)
            else:
                await update.message.reply_text("❌ Volume entre 0-100")
        except:
            await update.message.reply_text("❌ Valor inválido")
    
    elif comando in ['mute', 'silenciar']:
        result = VolumeControl.mute()
        await update.message.reply_text(result)
    
    elif comando in ['play', 'pause']:
        result = MediaControl.play_pause()
        await update.message.reply_text(result)
    
    elif comando in ['next', 'próxima', 'proxima']:
        result = MediaControl.next_track()
        await update.message.reply_text(result)
    
    elif comando in ['prev', 'anterior']:
        result = MediaControl.prev_track()
        await update.message.reply_text(result)
    
    elif comando in ['info', 'sistema']:
        result = SystemInfo.get_system_info()
        await update.message.reply_text(result)
    
    elif comando in ['matar', 'kill']:
        if not args:
            await update.message.reply_text("❌ Uso: matar [processo]")
            return
        result = TaskManager.kill_process_by_name(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['processos', 'tasks']:
        result = TaskManager.list_processes()
        await update.message.reply_text(result)
    
    elif comando in ['tempo', 'weather']:
        if not args:
            await update.message.reply_text("❌ Uso: tempo [cidade]\nEx: tempo Lisboa")
            return
        await update.message.reply_text("🌍 A procurar...")
        result = Weather.get_weather(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['word', 'documento']:
        if not args:
            await update.message.reply_text("❌ Uso: word [título] | [conteúdo]")
            return
        full_text = " ".join(args)
        if '|' not in full_text:
            await update.message.reply_text("❌ Use | para separar")
            return
        parts = full_text.split('|')
        result = WordAutomation.create_document(parts[0].strip(), parts[1].strip())
        await update.message.reply_text(result)
    
    elif comando in ['foto', 'webcam']:
        await update.message.reply_text("📸 A tirar foto...")
        filename = CameraControl.take_photo()
        if filename:
            await update.message.reply_photo(photo=open(filename, 'rb'))
            os.remove(filename)
        else:
            await update.message.reply_text("❌ Erro ao aceder webcam")

    elif comando in ['speedtest', 'velocidade']:
        await update.message.reply_text("🌐 A testar velocidade...")
        result = InternetControl.test_speed()
        await update.message.reply_text(result)
    
    elif comando in ['ip', 'rede']:
        result = InternetControl.get_ip_info()
        await update.message.reply_text(result)
    
    elif comando in ['download', 'descarregar']:
        if not args:
            await update.message.reply_text("❌ Uso: download [url]")
            return
        await update.message.reply_text("⬇️ A descarregar...")
        result = InternetControl.download_file(args[0], args[1] if len(args) > 1 else None)
        await update.message.reply_text(result)
    
    elif comando in ['ping', 'verificar']:
        if not args:
            await update.message.reply_text("❌ Uso: ping [site]")
            return
        result = InternetControl.check_website(" ".join(args))
        await update.message.reply_text(result)

    elif comando in ['timer', 'temporizador']:
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: timer [nome] [segundos]")
            return
        try:
            result = TimerControl.start_timer(args[0], int(args[1]))
            await update.message.reply_text(result)
        except:
            await update.message.reply_text("❌ Tempo inválido")
    
    elif comando in ['timers', 'cronometros']:
        result = TimerControl.list_timers()
        await update.message.reply_text(result)
    
    elif comando in ['verificar'] and len(args) > 0 and args[0] == 'timer':
        if len(args) < 2:
            await update.message.reply_text("❌ Uso: verificar timer [nome]")
            return
        result = TimerControl.check_timer(args[1])
        await update.message.reply_text(result)

    elif comando in ['bloquear', 'lock']:
        result = SecurityControl.lock_pc()
        await update.message.reply_text(result)

    elif comando in ['redimensionar', 'resize']:
        if len(args) < 3:
            await update.message.reply_text("❌ Uso: resize [img] [largura] [altura]")
            return
        try:
            result = ImageControl.resize_image(args[0], int(args[1]), int(args[2]))
            await update.message.reply_text(result)
        except:
            await update.message.reply_text("❌ Valores inválidos")
    
    elif comando in ['converter'] and len(args) > 1 and '.' in args[0]:
        result = ImageControl.convert_image(args[0], args[1])
        await update.message.reply_text(result)

    elif comando in ['falar', 'speak', 'tts']:
        if not args:
            await update.message.reply_text("❌ Uso: falar [texto]")
            return
        result = TTSControl.speak(" ".join(args))
        await update.message.reply_text(result)

    elif comando in ['notificacao', 'notify']:
        if not args:
            await update.message.reply_text("❌ Uso: notify [título] | [mensagem]")
            return
        full_text = " ".join(args)
        if '|' not in full_text:
            await update.message.reply_text("❌ Use | para separar")
            return
        parts = full_text.split('|', 1)
        result = NotificationControl.show_notification(parts[0].strip(), parts[1].strip())
        await update.message.reply_text(result)
    
    elif comando in ['limpar'] and len(args) > 0 and args[0] in ['temp', 'temporarios']:
        await update.message.reply_text("🧹 A limpar...")
        result = MaintenanceControl.clean_temp()
        await update.message.reply_text(result)
    
    elif comando in ['esvaziar', 'reciclagem']:
        result = MaintenanceControl.empty_recycle_bin()
        await update.message.reply_text(result)
    
    elif comando in ['disco', 'discos', 'espaco']:
        result = MaintenanceControl.check_disk_space()
        await update.message.reply_text(result)

    elif comando in ['wallpaper', 'fundo']:
        if not args:
            await update.message.reply_text("❌ Uso: wallpaper [caminho]")
            return
        result = CustomizationControl.change_wallpaper(" ".join(args))
        await update.message.reply_text(result)
    
    elif comando in ['gaming'] and (len(args) == 0 or args[0] in ['modo', 'mode']):
        result = GamingControl.gaming_mode()
        await update.message.reply_text(result)
    
    elif comando in ['boost', 'otimizar']:
        result = GamingControl.boost_performance()
        await update.message.reply_text(result)
    
    elif comando in ['ajuda', 'help', 'menu', 'comandos']:
        await update.message.reply_text("""🤖 **SISTEMA COMPLETO + WI-FI**

📶 **WI-FI (NOVO!):**
• `wifi` - ver rede atual
• `redes` ou `listar` - ver redes disponíveis
• `ligar [nome]` - ligar a rede (depois password)
• `desligar wifi` - desligar Wi-Fi
• `esquecer [nome]` - remover rede
• `guardadas` - ver redes guardadas

**🖥️ SISTEMA:**
• `abrir/fechar [app]` • `apps`
• `desligar` • `reiniciar` • `suspender`

**🖱️ AUTOMAÇÃO:**
• `rato [x] [y]` • `click` • `escrever [texto]`
• `tecla [nome]` • `atalho ctrl+c` • `screenshot`

**📁 FICHEIROS:**
• `procurar [nome]` • `copiar [origem] [dest]`
• `mover` • `apagar` • `renomear`
• `zip [nome.zip] [files]` • `unzip [zip]` • `info [file]`

**📝 NOTAS:**
• `nota [título] | [texto]` • `notas`
• `ver nota [nº]` • `apagar nota [nº]`

**🧮 CALC:**
• `calc 2+2*5` • `converter 10 km m`

**⏰ LEMBRETES:**
• `lembrete [texto] YYYY-MM-DD HH:MM`
• `lembretes` • `apagar lembrete [nº]`

**🌐 WEB:**
• `google [busca]` • `youtube [busca]`
• `site [url]` • `ping [site]`

**🎵 MULTIMÉDIA:**
• `volume [0-100]` • `mute` • `play` • `next` • `prev`

**📊 SISTEMA:**
• `info` • `processos` • `matar [processo]` • `disco`

**🌐 INTERNET:**
• `speedtest` • `ip` • `download [url]`

**⏱️ TIMER:**
• `timer [nome] [seg]` • `timers` • `verificar timer [nome]`

**🖼️ IMAGENS:**
• `resize [img] [w] [h]` • `converter [img] [formato]`

**📊 OUTROS:**
• `falar [texto]` • `notify [título] | [msg]`
• `tempo [cidade]` • `word [título] | [texto]`
• `foto` • `limpar temp` • `bloquear`
• `wallpaper [path]` • `gaming` • `boost`

**💬 Fala comigo! Estou pronto! 🚀**""")
    
    elif any(word in text_lower for word in ['olá', 'oi', 'hey', 'hello', 'bom dia']):
        await update.message.reply_text("Olá! 😊\nEscreve 'ajuda' para ver os comandos!\n\n✨ **Novo:** Controlo de Wi-Fi disponível!")
    
    elif any(word in text_lower for word in ['como estás', 'tudo bem']):
        await update.message.reply_text("Estou bem, pronto para executar comandos! 💪")
    
    elif any(word in text_lower for word in ['obrigado', 'thanks']):
        await update.message.reply_text("De nada! 😊")
    
    else:
        await update.message.reply_text("🤔 Não percebi. Escreve 'ajuda' para ver os comandos!")

def main():
    """Função principal"""
    print("🤖 Sistema de Controlo Total do PC + Wi-Fi via Telegram")
    print("=" * 60)
    
    if TELEGRAM_TOKEN == "SEU_TOKEN_AQUI":
        print("❌ ERRO: Configura o TELEGRAM_TOKEN primeiro!")
        return
    
    app = Application.builder().token(TELEGRAM_TOKEN).build()
    
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    
    def reminder_thread():
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        loop.run_until_complete(check_reminders(app))

    thread = threading.Thread(target=reminder_thread, daemon=True)
    thread.start()

    print("✅ Bot iniciado com controlo Wi-Fi!")
    print("📱 Envia 'ajuda' ao teu bot no Telegram")
    print("📶 Novidade: Controlo completo de Wi-Fi disponível!")
    print("🔄 A correr...")

    app.run_polling()  

class Utils:
    @staticmethod
    def google_search(query):
        """Abre o Google com a pesquisa"""
        try:
            import webbrowser
            search_url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
            webbrowser.open(search_url)
            return f"🔍 A pesquisar no Google: '{query}'"
        except Exception as e:
            return f"❌ Erro ao abrir o Google: {str(e)}"

    @staticmethod
    def youtube_search(query):
        """Abre o YouTube com a pesquisa"""
        try:
            import webbrowser
            search_url = f"https://www.youtube.com/results?search_query={query.replace(' ', '+')}"
            webbrowser.open(search_url)
            return f"🎥 A pesquisar no YouTube: '{query}'"
        except Exception as e:
            return f"❌ Erro ao abrir o YouTube: {str(e)}"

if __name__ == "__main__":
    main()  